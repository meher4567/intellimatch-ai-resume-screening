"""
Script to generate sample resume files in PDF and DOCX formats
This creates test data for the resume parser
"""

from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import os

# Get the directory of this script
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DIR = os.path.join(SCRIPT_DIR, 'sample_resumes')

def create_pdf_from_text(text_file, pdf_file):
    """Convert text file to PDF"""
    # Read text content
    with open(text_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create PDF
    doc = SimpleDocTemplate(pdf_file, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Split into lines and add to PDF
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            # Use different styles for headers (all caps lines)
            if line.isupper() and len(line) < 50:
                style = styles['Heading1']
            elif '|' in line and '@' in line:  # Contact info
                style = styles['Normal']
            else:
                style = styles['Normal']
            
            para = Paragraph(line, style)
            story.append(para)
        else:
            story.append(Spacer(1, 0.1*inch))
    
    doc.build(story)
    print(f"✓ Created PDF: {pdf_file}")

def create_docx_from_text(text_file, docx_file):
    """Convert text file to DOCX"""
    # Read text content
    with open(text_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create DOCX
    doc = Document()
    
    # Split into lines and add to DOCX
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            # Add heading for all-caps lines
            if line.isupper() and len(line) < 50:
                doc.add_heading(line, level=1)
            else:
                para = doc.add_paragraph(line)
                # Format contact info differently
                if '|' in line and '@' in line:
                    para.runs[0].font.size = Pt(10)
        else:
            doc.add_paragraph('')  # Empty line
    
    doc.save(docx_file)
    print(f"✓ Created DOCX: {docx_file}")

def main():
    """Generate all sample resume files"""
    print("Generating sample resume files...")
    print(f"Output directory: {SAMPLE_DIR}\n")
    
    # List of text files to convert
    text_files = [
        'john_doe_simple.txt',
        'sarah_chen_data_scientist.txt',
        'michael_rodriguez_pm.txt'
    ]
    
    for text_file in text_files:
        text_path = os.path.join(SAMPLE_DIR, text_file)
        
        if not os.path.exists(text_path):
            print(f"✗ Text file not found: {text_path}")
            continue
        
        # Generate PDF
        base_name = text_file.replace('.txt', '')
        pdf_path = os.path.join(SAMPLE_DIR, f"{base_name}.pdf")
        create_pdf_from_text(text_path, pdf_path)
        
        # Generate DOCX
        docx_path = os.path.join(SAMPLE_DIR, f"{base_name}.docx")
        create_docx_from_text(text_path, docx_path)
        
        print()
    
    print("✓ All sample files generated successfully!")
    print(f"\nTotal files created: {len(text_files) * 3} (txt, pdf, docx for each resume)")

if __name__ == '__main__':
    main()
