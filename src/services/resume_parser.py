"""
Resume Parser Service
Extracts text and metadata from PDF and DOCX resume files.

Enhanced with:
- Better text cleaning and normalization
- Robust field extraction with multiple patterns
- Education details with degree, major, year
- Work experience with company, title, duration
- Project extraction
- Certification extraction
- Language detection
- Salary and notice period extraction
- Layout detection and handling
"""
import re
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import Dict, Optional, List, Tuple
from datetime import datetime
import unicodedata


class ResumeParser:
    """Handles parsing of resume files (PDF, DOCX) to extract text and metadata."""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc'}
        self._compile_patterns()
    
    def _compile_patterns(self):
        """Precompile regex patterns for better performance."""
        # Email patterns (multiple formats)
        self.email_pattern = re.compile(
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b'
        )
        
        # Phone patterns (international)
        self.phone_patterns = [
            re.compile(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'),  # US
            re.compile(r'\+?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}'),      # International
        ]
        
        # Experience patterns
        self.experience_patterns = [
            re.compile(r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+experience', re.IGNORECASE),
            re.compile(r'experience[:\s]+(\d+)\+?\s*(?:years?|yrs?)', re.IGNORECASE),
        ]
        
        # Education degree patterns
        self.degree_patterns = {
            'phd': re.compile(r'\b(phd|ph\.?\s*d\.?|doctorate|doctoral)\b', re.IGNORECASE),
            'masters': re.compile(r'\b(master|masters|m\.?\s*s\.?|m\.?\s*a\.?|mba|mca|m\.tech|m\.sc)\b', re.IGNORECASE),
            'bachelors': re.compile(r'\b(bachelor|bachelors|b\.?\s*s\.?|b\.?\s*a\.?|b\.tech|b\.sc|b\.e\.?)\b', re.IGNORECASE),
            'associate': re.compile(r'\b(associate|a\.?\s*s\.?|a\.?\s*a\.?)\b', re.IGNORECASE),
        }
        
        # Date patterns for work experience
        self.date_patterns = [
            re.compile(r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]+\d{4}\b', re.IGNORECASE),
            re.compile(r'\b(0?[1-9]|1[0-2])/\d{4}\b'),  # MM/YYYY
            re.compile(r'\b\d{4}\b'),  # Just year
        ]
        
        # Salary patterns
        self.salary_patterns = [
            re.compile(r'(?:salary|compensation|ctc)[:\s]*(?:rs\.?|₹|inr)?\s*(\d+(?:,\d+)?(?:\.\d+)?)\s*(?:lpa|lacs?|lakhs?|k|thousand)?', re.IGNORECASE),
            re.compile(r'(?:expected|current)\s+(?:salary|ctc)[:\s]*(?:rs\.?|₹|inr)?\s*(\d+(?:,\d+)?(?:\.\d+)?)\s*(?:lpa|lacs?|lakhs?|k)?', re.IGNORECASE),
        ]
        
        # Notice period patterns
        self.notice_patterns = [
            re.compile(r'notice\s+period[:\s]*(\d+)\s*(days?|weeks?|months?)', re.IGNORECASE),
            re.compile(r'(?:can join|available)\s+(?:in|within)[:\s]*(\d+)\s*(days?|weeks?|months?)', re.IGNORECASE),
        ]
    
    def parse_file(self, file_path: str) -> Dict[str, any]:
        """
        Parse a resume file and extract structured data.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary with extracted data: text, name, email, phone, skills, etc.
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        file_ext = path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Extract text based on file type
        if file_ext == '.pdf':
            raw_text = self._extract_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            raw_text = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported format: {file_ext}")
        
        # Clean and normalize text
        cleaned_text = self._clean_text(raw_text)
        
        # Extract structured information
        parsed_data = {
            'raw_text': raw_text,
            'cleaned_text': cleaned_text,
            'email': self._extract_email(cleaned_text),
            'phone': self._extract_phone(cleaned_text),
            'name': self._extract_name(cleaned_text),
            'skills': self._extract_skills_basic(cleaned_text),
            'education': self._extract_education_detailed(cleaned_text),
            'experience': self._extract_work_experience(cleaned_text),
            'experience_years': self._estimate_experience(cleaned_text),
            'projects': self._extract_projects(cleaned_text),
            'certifications': self._extract_certifications(cleaned_text),
            'languages': self._extract_languages(cleaned_text),
            'salary_expectation': self._extract_salary(cleaned_text),
            'notice_period': self._extract_notice_period(cleaned_text),
            'linkedin': self._extract_linkedin(cleaned_text),
            'github': self._extract_github(cleaned_text),
            'parsed_at': datetime.utcnow().isoformat(),
            'quality_score': self._calculate_quality_score(raw_text, cleaned_text)
        }
        
        return parsed_data
    
    def _extract_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using PyMuPDF.
        Handles multiple extraction methods for different PDF types.
        """
        text = ""
        try:
            with fitz.open(file_path) as doc:
                for page_num, page in enumerate(doc):
                    # Try standard text extraction
                    page_text = page.get_text()
                    
                    # If page has little text, it might be scanned - try dict extraction
                    if len(page_text.strip()) < 50:
                        page_text = page.get_text("text")
                    
                    # Add page separator for multi-page resumes
                    if page_num > 0:
                        text += "\n\n--- PAGE BREAK ---\n\n"
                    
                    text += page_text
                    
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
        
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    
    def _extract_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx.
        Also extracts text from tables if present.
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract from tables (many resumes use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            text = "\n".join(text_parts)
            
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
        
        return text.strip()
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        - Remove excessive whitespace
        - Normalize unicode characters
        - Remove special characters that interfere with parsing
        """
        if not text:
            return ""
        
        # Normalize unicode (handles accented characters)
        text = unicodedata.normalize('NFKD', text)
        
        # Remove non-printable characters
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        
        # Normalize whitespace
        text = re.sub(r' +', ' ', text)  # Multiple spaces to single
        text = re.sub(r'\n +', '\n', text)  # Remove leading spaces on lines
        text = re.sub(r' +\n', '\n', text)  # Remove trailing spaces on lines
        text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 newlines
        
        return text.strip()
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text using precompiled regex."""
        match = self.email_pattern.search(text)
        return match.group(0) if match else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text using multiple patterns."""
        for pattern in self.phone_patterns:
            match = pattern.search(text)
            if match:
                phone = match.group(0)
                # Clean up the phone number
                phone = re.sub(r'[^\d+]', '', phone)
                if len(phone) >= 10:  # Valid phone should have at least 10 digits
                    return match.group(0)  # Return original format
        return None
    
    def _extract_name(self, text: str) -> Optional[str]:
        """
        Extract candidate name from resume.
        Assumes name is in the first few lines and is capitalized.
        """
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            # Look for lines with 2-4 words, each capitalized, no special chars
            if line and 2 <= len(line.split()) <= 4:
                if all(word[0].isupper() for word in line.split() if word):
                    # Skip if it looks like a company or title
                    if not any(keyword in line.lower() for keyword in ['resume', 'cv', 'curriculum', 'inc', 'llc', 'ltd']):
                        return line
        return None
    
    def _extract_skills_basic(self, text: str) -> List[str]:
        """
        Basic skill extraction using keyword matching.
        This will be enhanced with NLP in the next phase.
        """
        # Common technical skills (expand this list)
        skill_keywords = {
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'go', 'rust',
            'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'spring', 'fastapi',
            'sql', 'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'terraform',
            'machine learning', 'deep learning', 'nlp', 'computer vision', 'data science',
            'agile', 'scrum', 'git', 'ci/cd', 'rest api', 'graphql',
            'html', 'css', 'sass', 'webpack', 'babel',
            'pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras'
        }
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in skill_keywords:
            if skill in text_lower:
                found_skills.append(skill.title())
        
        return list(set(found_skills))  # Remove duplicates
    
    def _extract_education(self, text: str) -> List[str]:
        """Extract education information (degrees) - simple version for backward compatibility."""
        degrees = []
        for degree_level, pattern in self.degree_patterns.items():
            if pattern.search(text):
                degrees.append(degree_level)
        return degrees
    
    def _extract_education_detailed(self, text: str) -> List[Dict[str, any]]:
        """
        Extract detailed education information including degree, major, university, year.
        Returns list of education entries.
        """
        education_entries = []
        lines = text.split('\n')
        
        # Find education section
        education_start = -1
        for i, line in enumerate(lines):
            if re.search(r'\b(education|academic|qualification|degree)\b', line, re.IGNORECASE):
                education_start = i
                break
        
        if education_start == -1:
            # No explicit section, look for degree keywords
            for degree_level, pattern in self.degree_patterns.items():
                matches = list(pattern.finditer(text))
                for match in matches:
                    # Get context around the match (3 lines)
                    start = max(0, match.start() - 150)
                    end = min(len(text), match.end() + 150)
                    context = text[start:end]
                    
                    entry = {
                        'degree': degree_level,
                        'major': self._extract_major(context),
                        'university': self._extract_university(context),
                        'year': self._extract_year(context)
                    }
                    education_entries.append(entry)
        else:
            # Parse education section (next 10-15 lines after header)
            section_text = '\n'.join(lines[education_start:education_start + 15])
            
            for degree_level, pattern in self.degree_patterns.items():
                if pattern.search(section_text):
                    entry = {
                        'degree': degree_level,
                        'major': self._extract_major(section_text),
                        'university': self._extract_university(section_text),
                        'year': self._extract_year(section_text)
                    }
                    education_entries.append(entry)
        
        return education_entries
    
    def _extract_major(self, text: str) -> Optional[str]:
        """Extract major/field of study from education text."""
        # Common patterns: "in Computer Science", "B.S. Computer Science"
        major_patterns = [
            r'(?:in|of|major)\s+([A-Z][A-Za-z\s&]+?)(?:\s+from|\s+at|\s*,|\s*\n|$)',
            r'(?:B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?)\s+([A-Z][A-Za-z\s&]+?)(?:\s+from|\s+at|\s*,|\s*\n|$)'
        ]
        
        for pattern in major_patterns:
            match = re.search(pattern, text)
            if match:
                major = match.group(1).strip()
                if len(major) > 3 and len(major) < 50:  # Reasonable length
                    return major
        return None
    
    def _extract_university(self, text: str) -> Optional[str]:
        """Extract university/institution name from education text."""
        # Patterns: "from University of", "at MIT", "University Name"
        uni_patterns = [
            r'(?:from|at)\s+([A-Z][A-Za-z\s,&]+(?:University|College|Institute|School))',
            r'([A-Z][A-Za-z\s,&]+(?:University|College|Institute|School))'
        ]
        
        for pattern in uni_patterns:
            match = re.search(pattern, text)
            if match:
                uni = match.group(1).strip()
                if len(uni) > 5 and len(uni) < 100:
                    return uni
        return None
    
    def _extract_year(self, text: str) -> Optional[int]:
        """Extract graduation year from education text."""
        # Look for years (typically 1950-2030)
        year_pattern = r'\b(19[5-9]\d|20[0-3]\d)\b'
        matches = re.findall(year_pattern, text)
        if matches:
            # Return the most recent year (likely graduation)
            return int(max(matches))
        return None
    
    def _extract_work_experience(self, text: str) -> List[Dict[str, any]]:
        """
        Extract work experience entries with company, title, duration.
        """
        experiences = []
        lines = text.split('\n')
        
        # Find work experience section
        exp_start = -1
        for i, line in enumerate(lines):
            if re.search(r'\b(experience|employment|work history|professional)\b', line, re.IGNORECASE):
                exp_start = i
                break
        
        if exp_start == -1:
            return experiences  # No explicit experience section found
        
        # Parse experience section (typically 20-30 lines)
        section_lines = lines[exp_start:exp_start + 30]
        
        # Look for job entries (company/title patterns)
        current_entry = {}
        for line in section_lines:
            line = line.strip()
            if not line:
                # Empty line might indicate end of entry
                if current_entry:
                    experiences.append(current_entry)
                    current_entry = {}
                continue
            
            # Check if line has dates (likely a job entry)
            date_match = self.date_patterns[0].search(line) or self.date_patterns[1].search(line)
            if date_match:
                if current_entry:
                    experiences.append(current_entry)
                
                current_entry = {
                    'title': None,
                    'company': None,
                    'duration': line,
                    'description': []
                }
            elif current_entry:
                # Add to description or extract title/company
                if not current_entry['title'] and self._looks_like_job_title(line):
                    current_entry['title'] = line
                elif not current_entry['company']:
                    current_entry['company'] = line
                else:
                    current_entry['description'].append(line)
        
        # Add last entry
        if current_entry:
            experiences.append(current_entry)
        
        # Clean up descriptions
        for exp in experiences:
            if isinstance(exp.get('description'), list):
                exp['description'] = ' '.join(exp['description'][:3])  # First 3 lines only
        
        return experiences[:5]  # Return top 5 experiences
    
    def _looks_like_job_title(self, text: str) -> bool:
        """Check if text looks like a job title."""
        title_keywords = [
            'engineer', 'developer', 'manager', 'analyst', 'consultant',
            'director', 'lead', 'senior', 'junior', 'intern', 'architect',
            'designer', 'specialist', 'coordinator', 'administrator'
        ]
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in title_keywords)
    
    def _estimate_experience(self, text: str) -> Optional[int]:
        """
        Estimate years of experience from resume text.
        Looks for patterns like "5 years of experience" or date ranges.
        """
        # Pattern: "X years of experience"
        for pattern in self.experience_patterns:
            match = pattern.search(text)
            if match:
                return int(match.group(1))
        
        # Fallback: Count year ranges from work experience (rough estimate)
        year_pattern = r'\b(19[5-9]\d|20[0-3]\d)\b'
        years = re.findall(year_pattern, text)
        if len(years) >= 2:
            years = [int(y) for y in years]
            return max(years) - min(years)
        
        return None
    
    def _extract_projects(self, text: str) -> List[Dict[str, str]]:
        """Extract project information from resume."""
        projects = []
        lines = text.split('\n')
        
        # Find projects section
        proj_start = -1
        for i, line in enumerate(lines):
            if re.search(r'\b(projects?|portfolio)\b', line, re.IGNORECASE):
                proj_start = i
                break
        
        if proj_start == -1:
            return projects
        
        # Parse projects (next 15 lines)
        section_lines = lines[proj_start:proj_start + 15]
        current_project = {}
        
        for line in section_lines:
            line = line.strip()
            if not line:
                if current_project:
                    projects.append(current_project)
                    current_project = {}
                continue
            
            # First non-empty line is likely project name
            if not current_project:
                current_project = {'name': line, 'description': []}
            else:
                current_project['description'].append(line)
        
        if current_project:
            projects.append(current_project)
        
        # Clean up descriptions
        for proj in projects:
            if isinstance(proj.get('description'), list):
                proj['description'] = ' '.join(proj['description'][:2])
        
        return projects[:3]  # Top 3 projects
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications from resume."""
        certifications = []
        
        # Common certification keywords
        cert_keywords = [
            'aws certified', 'azure certified', 'gcp certified',
            'pmp', 'cissp', 'comptia', 'cisco', 'oracle certified',
            'certified scrum', 'csm', 'cka', 'ckad',
            'tensorflow', 'microsoft certified'
        ]
        
        text_lower = text.lower()
        for cert in cert_keywords:
            if cert in text_lower:
                # Get the line containing the certification
                for line in text.split('\n'):
                    if cert in line.lower():
                        certifications.append(line.strip())
                        break
        
        return list(set(certifications))[:5]  # Top 5 unique certs
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extract spoken languages from resume."""
        languages = []
        
        # Common language names
        lang_keywords = [
            'english', 'spanish', 'french', 'german', 'chinese', 'japanese',
            'hindi', 'arabic', 'portuguese', 'russian', 'italian', 'korean',
            'tamil', 'telugu', 'bengali', 'marathi', 'gujarati'
        ]
        
        # Look for language section
        lines = text.split('\n')
        for i, line in enumerate(lines):
            if re.search(r'\b(languages?|linguistic)\b', line, re.IGNORECASE):
                # Check next 3 lines
                for j in range(i, min(i + 4, len(lines))):
                    line_lower = lines[j].lower()
                    for lang in lang_keywords:
                        if lang in line_lower:
                            languages.append(lang.title())
        
        # If no section found, do basic search
        if not languages:
            text_lower = text.lower()
            for lang in lang_keywords:
                if lang in text_lower:
                    languages.append(lang.title())
        
        return list(set(languages))
    
    def _extract_salary(self, text: str) -> Optional[str]:
        """Extract salary expectation from resume."""
        for pattern in self.salary_patterns:
            match = pattern.search(text)
            if match:
                return match.group(0)
        return None
    
    def _extract_notice_period(self, text: str) -> Optional[str]:
        """Extract notice period from resume."""
        for pattern in self.notice_patterns:
            match = pattern.search(text)
            if match:
                return match.group(0)
        return None
    
    def _extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn profile URL."""
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
        match = re.search(linkedin_pattern, text, re.IGNORECASE)
        return match.group(0) if match else None
    
    def _extract_github(self, text: str) -> Optional[str]:
        """Extract GitHub profile URL."""
        github_pattern = r'(?:https?://)?(?:www\.)?github\.com/[\w-]+'
        match = re.search(github_pattern, text, re.IGNORECASE)
        return match.group(0) if match else None
    
    def _calculate_quality_score(self, raw_text: str, cleaned_text: str) -> float:
        """
        Calculate resume quality score (0-100) based on:
        - Length (longer is better, up to a point)
        - Structure (sections, formatting)
        - Completeness (has email, phone, etc.)
        """
        score = 0.0
        
        # Length score (max 30 points)
        text_length = len(cleaned_text)
        if text_length > 2000:
            score += 30
        elif text_length > 1000:
            score += 20
        elif text_length > 500:
            score += 10
        
        # Structure score (max 40 points)
        structure_keywords = [
            'education', 'experience', 'skills', 'projects',
            'certification', 'work', 'employment'
        ]
        sections_found = sum(1 for keyword in structure_keywords 
                           if re.search(rf'\b{keyword}\b', cleaned_text, re.IGNORECASE))
        score += min(sections_found * 8, 40)
        
        # Completeness score (max 30 points)
        if self._extract_email(cleaned_text):
            score += 10
        if self._extract_phone(cleaned_text):
            score += 10
        if len(self._extract_skills_basic(cleaned_text)) > 3:
            score += 10
        
        return round(min(score, 100), 2)


# Utility functions
def validate_resume_file(file_path: str) -> bool:
    """Check if file exists and has a supported format."""
    parser = ResumeParser()
    path = Path(file_path)
    return path.exists() and path.suffix.lower() in parser.supported_formats


def get_file_metadata(file_path: str) -> Dict[str, any]:
    """Get file metadata (size, creation date, etc.)."""
    path = Path(file_path)
    if not path.exists():
        return {}
    
    stat = path.stat()
    return {
        'file_name': path.name,
        'file_size': stat.st_size,
        'file_extension': path.suffix,
        'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
        'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat()
    }
