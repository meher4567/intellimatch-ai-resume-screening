"""
DOCX Text Extractor
Extracts text content from Microsoft Word (.docx) files
"""

from docx import Document
from pathlib import Path
from typing import Dict, List
import logging

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extract text from DOCX files"""
    
    def __init__(self):
        """Initialize DOCX extractor"""
        self.supported_extensions = ['.docx']
    
    def extract_text(self, file_path: str) -> Dict[str, str]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dict with 'text' and 'method' keys
        """
        file_path = str(Path(file_path).resolve())
        
        # Validate file exists
        if not Path(file_path).exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        # Validate file extension
        if Path(file_path).suffix.lower() not in self.supported_extensions:
            raise ValueError(f"Not a DOCX file: {file_path}")
        
        logger.info(f"Extracting text from DOCX: {file_path}")
        
        try:
            # Open document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_texts.append(" | ".join(row_text))
            
            # Combine all text
            all_text = paragraphs + table_texts
            full_text = "\n".join(all_text)
            
            logger.info(f"DOCX: Extracted {len(full_text)} characters from {len(paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            return {
                "text": full_text,
                "method": "python-docx",
                "num_paragraphs": len(paragraphs),
                "num_tables": len(doc.tables)
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return {"text": "", "method": "python-docx"}
    
    def extract_with_formatting(self, file_path: str) -> List[Dict]:
        """
        Extract text with formatting information (bold, italic, font size, etc.)
        Useful for identifying section headers
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of dicts with text and formatting info
        """
        file_path = str(Path(file_path).resolve())
        
        try:
            doc = Document(file_path)
            formatted_text = []
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                # Get paragraph style
                style_name = para.style.name if para.style else "Normal"
                
                # Check if it's a heading
                is_heading = style_name.startswith("Heading")
                
                # Extract runs (text with consistent formatting)
                runs_info = []
                for run in para.runs:
                    if run.text.strip():
                        runs_info.append({
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline,
                            "font_size": run.font.size.pt if run.font.size else None
                        })
                
                formatted_text.append({
                    "text": para.text,
                    "style": style_name,
                    "is_heading": is_heading,
                    "runs": runs_info
                })
            
            logger.info(f"Extracted {len(formatted_text)} formatted paragraphs")
            return formatted_text
            
        except Exception as e:
            logger.error(f"Failed to extract formatted text from {file_path}: {e}")
            return []
    
    def get_document_properties(self, file_path: str) -> Dict:
        """
        Extract document properties/metadata
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dict with document properties
        """
        try:
            doc = Document(file_path)
            props = doc.core_properties
            
            metadata = {
                "author": props.author or "",
                "title": props.title or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
            }
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract properties from {file_path}: {e}")
            return {}


# Convenience function
def extract_text_from_docx(file_path: str) -> str:
    """
    Convenience function to extract text from DOCX
    
    Args:
        file_path: Path to DOCX file
    
    Returns:
        Extracted text content
    """
    extractor = DOCXExtractor()
    result = extractor.extract_text(file_path)
    return result["text"]
